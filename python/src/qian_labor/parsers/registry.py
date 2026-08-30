from __future__ import annotations

import csv
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any

import pymupdf as fitz
import xlrd
from charset_normalizer import from_bytes
from docx import Document
from openpyxl import load_workbook
from PIL import Image, ImageOps

from qian_labor.parsers.protocols import ParsedBlock, ParsedDocument, VisionPage

MAX_ROWS = 10_000
MAX_COLUMNS = 200
MAX_PDF_PAGES = 100
MAX_RENDER_PIXELS = 8_000_000


class ParserRegistry:
    def choose(self, filename: str, content: bytes) -> str:
        extension = Path(filename).suffix.lower()
        if extension in {".xlsx", ".xls", ".csv"}:
            return "spreadsheet"
        if extension == ".docx":
            return "docx"
        if extension == ".pdf":
            return "pdf"
        if extension in {".jpg", ".jpeg", ".png", ".webp"}:
            return "image"
        raise ValueError("unsupported parser")

    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        extension = Path(filename).suffix.lower()
        if extension == ".csv":
            return self._parse_csv(content)
        if extension == ".xlsx":
            return self._parse_xlsx(content)
        if extension == ".xls":
            return self._parse_xls(content)
        if extension == ".docx":
            return self._parse_docx(content)
        if extension == ".pdf":
            return self._parse_pdf(content)
        if extension in {".jpg", ".jpeg", ".png", ".webp"}:
            return self._parse_image(content, extension)
        raise ValueError("unsupported parser")

    def _parse_csv(self, content: bytes) -> ParsedDocument:
        match = from_bytes(content).best()
        if match is None:
            raise ValueError("CSV_ENCODING_UNKNOWN")
        text = str(match)
        rows = list(csv.reader(StringIO(text)))
        if not rows:
            return ParsedDocument("spreadsheet", warnings=["empty_csv"])
        if len(rows) > MAX_ROWS or max(map(len, rows), default=0) > MAX_COLUMNS:
            raise ValueError("CSV_DIMENSION_LIMIT")
        headers = [value.strip() for value in rows[0]]
        blocks = [
            ParsedBlock(
                text=" | ".join(headers),
                block_type="header",
                locator={"sheet": "CSV", "row": 1, "headers": headers},
            )
        ]
        for row_number, row in enumerate(rows[1:], start=2):
            blocks.append(
                ParsedBlock(
                    text=" | ".join(row),
                    block_type="row",
                    locator={
                        "sheet": "CSV",
                        "row": row_number,
                        "columns": len(row),
                        "headers": headers[: len(row)],
                    },
                )
            )
        return ParsedDocument("spreadsheet", blocks)

    def _parse_xlsx(self, content: bytes) -> ParsedDocument:
        workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
        blocks: list[ParsedBlock] = []
        for sheet in workbook.worksheets:
            headers = [
                "" if cell.value is None else str(cell.value)
                for cell in next(sheet.iter_rows(min_row=1, max_row=1), ())
            ][:MAX_COLUMNS]
            blocks.append(
                ParsedBlock(
                    text=" | ".join(headers),
                    block_type="header",
                    locator={"sheet": sheet.title, "row": 1, "headers": headers},
                )
            )
            for row in sheet.iter_rows(min_row=2, max_row=MAX_ROWS, max_col=MAX_COLUMNS):
                for cell in row:
                    if cell.value is None:
                        continue
                    blocks.append(
                        ParsedBlock(
                            text=str(cell.value),
                            block_type="cell",
                            locator={
                                "sheet": sheet.title,
                                "row": cell.row,
                                "column": cell.column,
                                "cell": cell.coordinate,
                                "header": headers[cell.column - 1]
                                if cell.column <= len(headers)
                                else "",
                            },
                        )
                    )
        return ParsedDocument("spreadsheet", blocks)

    def _parse_xls(self, content: bytes) -> ParsedDocument:
        workbook = xlrd.open_workbook(file_contents=content, on_demand=True)
        blocks: list[ParsedBlock] = []
        for sheet in workbook.sheets():
            headers = [str(sheet.cell_value(0, column)) for column in range(sheet.ncols)]
            blocks.append(
                ParsedBlock(
                    text=" | ".join(headers),
                    block_type="header",
                    locator={"sheet": sheet.name, "row": 1, "headers": headers},
                )
            )
            for row in range(1, min(sheet.nrows, MAX_ROWS)):
                for column in range(min(sheet.ncols, MAX_COLUMNS)):
                    value = sheet.cell_value(row, column)
                    if value == "":
                        continue
                    blocks.append(
                        ParsedBlock(
                            text=str(value),
                            block_type="cell",
                            locator={
                                "sheet": sheet.name,
                                "row": row + 1,
                                "column": column + 1,
                                "cell": f"{xlrd.formula.colname(column)}{row + 1}",
                                "header": headers[column],
                            },
                        )
                    )
        return ParsedDocument("spreadsheet", blocks)

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        document = Document(BytesIO(content))
        blocks: list[ParsedBlock] = []
        for index, paragraph in enumerate(document.paragraphs, start=1):
            if paragraph.text.strip():
                blocks.append(
                    ParsedBlock(
                        text=paragraph.text.strip(),
                        block_type="paragraph",
                        locator={"paragraph": index},
                    )
                )
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for column_index, cell in enumerate(row.cells, start=1):
                    if cell.text.strip():
                        blocks.append(
                            ParsedBlock(
                                text=cell.text.strip(),
                                block_type="table_cell",
                                locator={
                                    "table": table_index,
                                    "row": row_index,
                                    "column": column_index,
                                },
                            )
                        )
        warnings = ["embedded_images_need_vision"] if document.inline_shapes else []
        return ParsedDocument("docx", blocks, warnings=warnings)

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        document = fitz.open(stream=content, filetype="pdf")
        if document.page_count > MAX_PDF_PAGES:
            raise ValueError("PDF_PAGE_LIMIT")
        blocks: list[ParsedBlock] = []
        vision_pages: list[VisionPage] = []
        pages: Any = document
        for page_number, page in enumerate(pages, start=1):
            page_blocks = page.get_text("blocks")
            useful_text = "".join(str(item[4]).strip() for item in page_blocks)
            if len(useful_text) >= 8:
                for item in page_blocks:
                    text = str(item[4]).strip()
                    if not text:
                        continue
                    blocks.append(
                        ParsedBlock(
                            text=text,
                            block_type="pdf_text",
                            locator={
                                "page": page_number,
                                "bbox": [round(float(value), 2) for value in item[:4]],
                            },
                        )
                    )
            else:
                matrix = fitz.Matrix(1.5, 1.5)
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                if pixmap.width * pixmap.height > MAX_RENDER_PIXELS:
                    raise ValueError("PDF_RENDER_PIXEL_LIMIT")
                vision_pages.append(
                    VisionPage(
                        page=page_number,
                        media_type="image/png",
                        image_bytes=pixmap.tobytes("png"),
                        width=pixmap.width,
                        height=pixmap.height,
                    )
                )
        document.close()
        return ParsedDocument(
            "pdf",
            blocks,
            needs_vision=bool(vision_pages),
            vision_pages=vision_pages,
        )

    def _parse_image(self, content: bytes, extension: str) -> ParsedDocument:
        with Image.open(BytesIO(content)) as raw:
            image = ImageOps.exif_transpose(raw)
            width, height = image.size
            output = BytesIO()
            image.convert("RGB").save(output, format="JPEG", quality=88)
        return ParsedDocument(
            "image",
            needs_vision=True,
            vision_pages=[
                VisionPage(
                    page=1,
                    media_type="image/jpeg",
                    image_bytes=output.getvalue(),
                    width=width,
                    height=height,
                )
            ],
        )
