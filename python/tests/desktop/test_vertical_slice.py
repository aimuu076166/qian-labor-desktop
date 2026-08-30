from __future__ import annotations

import json
import time
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from qian_labor.desktop.app import create_desktop_app

TOKEN = "synthetic-vertical-token"
HEADERS = {"X-Qian-Desktop-Token": TOKEN}
TERMINAL = {"completed", "matching_review", "partial", "failed"}


def _write_synthetic_contract(path: Path) -> None:
    payload = {
        "synthetic_marker": "QIAN_DEMO_20260824",
        "document_type": "contract",
        "employee_number": "F-701",
        "employee_name": "完全虚构员工甲",
        "department": "虚构制造部",
        "job_title": "虚构操作员",
        "facts": {
            "employment.status": "active",
            "employment.contract.exists": False,
            "employment.identity.match_status": "confirmed",
            "employment.material_coverage": 0.4,
            "analysis.minimum_core_coverage": 0.4,
        },
    }
    document = Document()
    document.add_heading("完全虚构劳动合同测试材料", level=1)
    document.add_paragraph(
        "QIAN_SYNTHETIC_JSON=" + json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
    )
    document.save(path)


def _poll_terminal(client: TestClient, analysis_id: str) -> dict[str, object]:
    deadline = time.monotonic() + 10
    while time.monotonic() < deadline:
        response = client.get(
            f"/api/analyses/{analysis_id}/processing",
            headers=HEADERS,
        )
        assert response.status_code == 200
        payload = response.json()
        if payload["status"] in TERMINAL:
            return payload
        time.sleep(0.02)
    raise AssertionError("desktop synthetic analysis did not reach a terminal state")


def test_synthetic_desktop_vertical_slice_is_traceable_deleted_and_persistent(tmp_path: Path) -> None:
    data_dir = tmp_path / "app-data"
    source = tmp_path / "fictional-contract.docx"
    _write_synthetic_contract(source)

    app = create_desktop_app(data_dir=data_dir, launch_token=TOKEN)
    with TestClient(app) as client:
        created = client.post(
            "/api/analyses",
            headers=HEADERS,
            json={"name": "虚构企业一键体检", "company_display_name": "完全虚构企业"},
        )
        assert created.status_code == 201
        analysis_id = created.json()["id"]

        imported = client.post(
            f"/api/analyses/{analysis_id}/import-paths",
            headers=HEADERS,
            json={"paths": [str(source)]},
        )
        assert imported.status_code == 200

        submitted = client.post(f"/api/analyses/{analysis_id}/process", headers=HEADERS)
        assert submitted.status_code == 202
        terminal = _poll_terminal(client, analysis_id)
        assert terminal["status"] == "completed"

        dashboard = client.get(f"/api/analyses/{analysis_id}/dashboard", headers=HEADERS)
        assert dashboard.status_code == 200
        dashboard_payload = dashboard.json()
        assert dashboard_payload["summary"]["employee_count"] == 1
        assert dashboard_payload["summary"]["finding_count"] >= 1
        assert dashboard_payload["summary"]["insufficient_data_count"] >= 1
        assert dashboard_payload["findings"]

        finding_id = dashboard_payload["findings"][0]["id"]
        finding = client.get(f"/api/findings/{finding_id}", headers=HEADERS)
        assert finding.status_code == 200
        finding_payload = finding.json()
        assert finding_payload["rule_id"]
        assert finding_payload["assessment_status"] != "not_triggered"
        assert finding_payload["sources"]
        assert finding_payload["sources"][0]["file_name"] == "fictional-contract.docx"
        assert finding_payload["sources"][0]["location"]

        deleted = client.delete(f"/api/analyses/{analysis_id}", headers=HEADERS)
        assert deleted.status_code == 200
        assert deleted.json()["status"] == "deleted"
        assert client.get(f"/api/analyses/{analysis_id}/dashboard", headers=HEADERS).status_code == 404
        assert client.get(f"/api/findings/{finding_id}", headers=HEADERS).status_code == 404

    # Re-open the sidecar against the exact same SQLite file: deletion must remain durable.
    restarted = create_desktop_app(data_dir=data_dir, launch_token=TOKEN)
    with TestClient(restarted) as client:
        assert client.get(f"/api/analyses/{analysis_id}/dashboard", headers=HEADERS).status_code == 404
        assert client.get(f"/api/findings/{finding_id}", headers=HEADERS).status_code == 404
