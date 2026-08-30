from io import BytesIO

import fitz
import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from qian_labor.parsers.registry import ParserRegistry
from qian_labor.security.uploads import UploadRejected, validate_upload


def xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "花名册"
    sheet.append(["员工编号", "姓名", "部门"])
    sheet.append(["S-001", "虚构甲", "制造部"])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("劳动合同（完全虚构）")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "员工编号"
    table.cell(0, 1).text = "S-001"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_bytes(text: str = "完全虚构劳动合同 S-001") -> bytes:
    document = fitz.open()
    page = document.new_page()
    if text:
        page.insert_text((72, 72), text)
    value = document.tobytes()
    document.close()
    return value


def png_bytes() -> bytes:
    output = BytesIO()
    Image.new("RGB", (80, 40), "white").save(output, format="PNG")
    return output.getvalue()


def test_csv_and_xlsx_preserve_row_and_cell_locations() -> None:
    csv_document = ParserRegistry().parse("roster.csv", "员工编号,姓名\nS-001,虚构甲\n".encode())
    xlsx_document = ParserRegistry().parse("roster.xlsx", xlsx_bytes())

    assert csv_document.blocks[1].locator["row"] == 2
    assert xlsx_document.blocks[1].locator == {
        "sheet": "花名册",
        "row": 2,
        "column": 1,
        "cell": "A2",
        "header": "员工编号",
    }


def test_docx_and_text_pdf_preserve_structural_locations() -> None:
    word = ParserRegistry().parse("contract.docx", docx_bytes())
    pdf = ParserRegistry().parse("contract.pdf", pdf_bytes())

    assert word.blocks[0].locator["paragraph"] == 1
    assert word.blocks[1].locator["table"] == 1
    assert pdf.blocks[0].locator["page"] == 1
    assert len(pdf.blocks[0].locator["bbox"]) == 4


def test_scanned_pdf_and_image_require_vision_without_fabricated_text() -> None:
    scan = ParserRegistry().parse("scan.pdf", pdf_bytes(""))
    image = ParserRegistry().parse("photo.png", png_bytes())

    assert scan.needs_vision and scan.blocks == [] and scan.vision_pages[0].page == 1
    assert image.needs_vision and image.blocks == []


def test_rejects_executable_disguised_as_pdf() -> None:
    with pytest.raises(UploadRejected):
        validate_upload("note.pdf", "application/pdf", b"MZ fake")


def test_accepts_synthetic_png() -> None:
    output = BytesIO()
    Image.new("RGB", (20, 20), "white").save(output, format="PNG")
    assert len(validate_upload("demo.png", "image/png", output.getvalue())) == 64


def test_rejects_path_traversal_and_corrupt_images() -> None:
    with pytest.raises(UploadRejected):
        validate_upload("../demo.csv", "text/csv", b"name\nsynthetic\n")
    with pytest.raises(UploadRejected):
        validate_upload("demo.png", "image/png", b"\x89PNG\r\n\x1a\ncorrupt")
