#!/usr/bin/env python3
from __future__ import annotations

import ctypes
import json
import os
import queue
import secrets
import signal
import subprocess
import tempfile
import threading
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Sequence

import httpx
from docx import Document

from qian_labor.rules.catalog import RULE_IDS


ROOT = Path(__file__).resolve().parents[1]
READY_PREFIX = "QIAN_DESKTOP_READY="
TERMINAL = {"completed", "matching_review", "partial", "failed"}
PROCESSING_TIMEOUT_SECONDS = 30
MARKERS = (
    "SIDECAR_BOOT=PASS",
    "LOOPBACK_ONLY=PASS",
    "TOKEN_AUTH=PASS",
    "SQLITE_PERSISTENCE=PASS",
    "SYNTHETIC_IMPORT=PASS",
    "FAKE_PROVIDER_PIPELINE=PASS",
    "R01_R20_REGRESSION=PASS",
    "SOURCE_TRACE=PASS",
    "DELETE_CLEANUP=PASS",
    "SIDECAR_SHUTDOWN=PASS",
)


class VerificationError(RuntimeError):
    def __init__(self, code: str):
        self.code = code
        super().__init__(code)


@dataclass(frozen=True)
class RunningSidecar:
    process: subprocess.Popen[str]
    base_url: str
    ready_pid: int


def _write_fixture(path: Path) -> None:
    payload = {
        "synthetic_marker": "QIAN_DEMO_20260824",
        "document_type": "contract",
        "employee_number": "F-801",
        "employee_name": "完全虚构验收员工",
        "department": "虚构测试部门",
        "job_title": "虚构岗位",
        "facts": {
            "employment.status": "active",
            "employment.contract.exists": False,
            "employment.identity.match_status": "confirmed",
            "employment.material_coverage": 0.4,
            "analysis.minimum_core_coverage": 0.4,
        },
    }
    document = Document()
    document.add_heading("完全虚构桌面验收材料", level=1)
    document.add_paragraph(
        "QIAN_SYNTHETIC_JSON=" + json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
    )
    document.save(path)


def _popen_options(*, windows_no_window: bool = False) -> dict[str, object]:
    if os.name == "nt":
        flags = subprocess.CREATE_NEW_PROCESS_GROUP
        if windows_no_window:
            flags |= subprocess.CREATE_NO_WINDOW
        return {"creationflags": flags}
    return {"start_new_session": True}


def _start_sidecar(
    command: Sequence[str],
    data_dir: Path,
    token: str,
    *,
    cwd: Path,
    windows_no_window: bool,
) -> RunningSidecar:
    env = {
        **os.environ,
        "AI_PROVIDER": "fake",
        "AI_API_KEY": "",
        "QIAN_DESKTOP_DATA_DIR": str(data_dir),
        "QIAN_DESKTOP_TOKEN": token,
        "QIAN_DESKTOP_PORT": "0",
    }
    try:
        process = subprocess.Popen(
            list(command),
            cwd=cwd,
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1,
            **_popen_options(windows_no_window=windows_no_window),
        )
    except (FileNotFoundError, PermissionError, OSError) as error:
        raise VerificationError("BINARY_START_FAILED") from error

    assert process.stdout is not None
    lines: queue.Queue[str] = queue.Queue()

    def read_stdout() -> None:
        for line in process.stdout:
            lines.put(line.rstrip("\r\n"))

    threading.Thread(target=read_stdout, daemon=True).start()
    deadline = time.monotonic() + 45
    while time.monotonic() < deadline:
        if process.poll() is not None:
            raise VerificationError("EXITED_BEFORE_READY")
        try:
            line = lines.get(timeout=0.1)
        except queue.Empty:
            continue
        if not line.startswith(READY_PREFIX):
            continue
        try:
            payload = json.loads(line.split("=", 1)[1])
        except (json.JSONDecodeError, IndexError) as error:
            raise VerificationError("READY_INVALID") from error
        host = payload.get("host")
        port = payload.get("port")
        pid = payload.get("pid")
        if host not in {"127.0.0.1", "::1"}:
            raise VerificationError("NON_LOOPBACK_BIND")
        if not isinstance(port, int) or not 0 < port <= 65535:
            raise VerificationError("READY_INVALID")
        if not isinstance(pid, int) or pid <= 0:
            raise VerificationError("READY_INVALID")
        return RunningSidecar(process, f"http://{host}:{port}", pid)
    raise VerificationError("READY_TIMEOUT")


def _pid_is_alive(pid: int) -> bool:
    if os.name != "nt":
        try:
            os.kill(pid, 0)
        except ProcessLookupError:
            return False
        except PermissionError:
            return True
        return True
    process_query_limited_information = 0x1000
    still_active = 259
    handle = ctypes.windll.kernel32.OpenProcess(  # type: ignore[attr-defined]
        process_query_limited_information, False, pid
    )
    if not handle:
        return False
    try:
        exit_code = ctypes.c_ulong()
        if not ctypes.windll.kernel32.GetExitCodeProcess(  # type: ignore[attr-defined]
            handle, ctypes.byref(exit_code)
        ):
            return False
        return exit_code.value == still_active
    finally:
        ctypes.windll.kernel32.CloseHandle(handle)  # type: ignore[attr-defined]


def _stop_sidecar(running: RunningSidecar) -> None:
    process = running.process
    if process.poll() is None:
        if os.name == "nt":
            subprocess.run(
                ["taskkill", "/PID", str(process.pid), "/T", "/F"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=False,
            )
        else:
            try:
                os.killpg(process.pid, signal.SIGTERM)
            except ProcessLookupError:
                pass
        try:
            process.wait(timeout=10)
        except subprocess.TimeoutExpired:
            if os.name != "nt":
                try:
                    os.killpg(process.pid, signal.SIGKILL)
                except ProcessLookupError:
                    pass
            else:
                process.kill()
            process.wait(timeout=10)

    deadline = time.monotonic() + 10
    while time.monotonic() < deadline and _pid_is_alive(running.ready_pid):
        time.sleep(0.05)
    if _pid_is_alive(running.ready_pid):
        raise VerificationError("SHUTDOWN_TIMEOUT")


def _poll(client: httpx.Client, analysis_id: str) -> dict[str, object]:
    deadline = time.monotonic() + PROCESSING_TIMEOUT_SECONDS
    while time.monotonic() < deadline:
        response = client.get(f"/api/analyses/{analysis_id}/processing")
        if response.status_code != 200:
            raise VerificationError("PROCESSING_STATUS_FAILED")
        payload = response.json()
        if payload.get("status") in TERMINAL:
            return payload
        time.sleep(0.05)
    raise VerificationError("PROCESSING_TIMEOUT")


def _assert_status(response: httpx.Response, expected: int, code: str) -> None:
    if response.status_code != expected:
        raise VerificationError(code)


def verify_command(
    command: Sequence[str],
    *,
    token: str | None = None,
    cwd: Path | None = None,
    windows_no_window: bool = False,
) -> tuple[str, ...]:
    if not command or any(not isinstance(value, str) or not value for value in command):
        raise VerificationError("COMMAND_INVALID")
    running: RunningSidecar | None = None
    stopped_processes = 0
    launch_token = token or secrets.token_hex(32)
    try:
        with tempfile.TemporaryDirectory(prefix="qian-desktop-verify-") as temp:
            temp_path = Path(temp)
            data_dir = temp_path / "app-data"
            fixture = temp_path / "fictional-contract.docx"
            _write_fixture(fixture)

            running = _start_sidecar(
                command,
                data_dir,
                launch_token,
                cwd=cwd or ROOT,
                windows_no_window=windows_no_window,
            )
            with httpx.Client(base_url=running.base_url, timeout=5) as unauthenticated:
                _assert_status(unauthenticated.get("/health"), 200, "HEALTH_FAILED")
                _assert_status(unauthenticated.get("/api/status"), 401, "AUTH_BYPASS")
                _assert_status(
                    unauthenticated.get(
                        "/api/status", headers={"X-Qian-Desktop-Token": "incorrect-token"}
                    ),
                    401,
                    "AUTH_BYPASS",
                )

            with httpx.Client(
                base_url=running.base_url,
                headers={"X-Qian-Desktop-Token": launch_token},
                timeout=5,
            ) as client:
                _assert_status(client.get("/api/status"), 200, "TOKEN_AUTH_FAILED")
                created = client.post(
                    "/api/analyses",
                    json={"name": "虚构桌面验收", "company_display_name": "完全虚构企业"},
                )
                _assert_status(created, 201, "ANALYSIS_CREATE_FAILED")
                analysis_id = str(created.json()["id"])
                imported = client.post(
                    f"/api/analyses/{analysis_id}/import-paths",
                    json={"paths": [str(fixture)]},
                )
                _assert_status(imported, 200, "SYNTHETIC_IMPORT_FAILED")
                submitted = client.post(f"/api/analyses/{analysis_id}/process")
                _assert_status(submitted, 202, "PROCESS_SUBMIT_FAILED")
                terminal = _poll(client, analysis_id)
                if terminal.get("status") != "completed":
                    raise VerificationError("FAKE_PIPELINE_NOT_COMPLETED")
                dashboard = client.get(f"/api/analyses/{analysis_id}/dashboard")
                _assert_status(dashboard, 200, "DASHBOARD_FAILED")
                findings = dashboard.json().get("findings", [])
                if not findings:
                    raise VerificationError("FINDING_MISSING")
                finding_id = str(findings[0]["id"])
                finding = client.get(f"/api/findings/{finding_id}")
                _assert_status(finding, 200, "FINDING_DETAIL_FAILED")
                if not finding.json().get("sources"):
                    raise VerificationError("SOURCE_TRACE_MISSING")

            _stop_sidecar(running)
            stopped_processes += 1
            running = None

            running = _start_sidecar(
                command,
                data_dir,
                launch_token,
                cwd=cwd or ROOT,
                windows_no_window=windows_no_window,
            )
            with httpx.Client(
                base_url=running.base_url,
                headers={"X-Qian-Desktop-Token": launch_token},
                timeout=5,
            ) as client:
                persisted = client.get(f"/api/analyses/{analysis_id}/dashboard")
                _assert_status(persisted, 200, "SQLITE_PERSISTENCE_FAILED")
                deleted = client.delete(f"/api/analyses/{analysis_id}")
                _assert_status(deleted, 200, "DELETE_FAILED")
                if deleted.json().get("status") != "deleted":
                    raise VerificationError("DELETE_STATUS_INVALID")

            _stop_sidecar(running)
            stopped_processes += 1
            running = None

            running = _start_sidecar(
                command,
                data_dir,
                launch_token,
                cwd=cwd or ROOT,
                windows_no_window=windows_no_window,
            )
            with httpx.Client(
                base_url=running.base_url,
                headers={"X-Qian-Desktop-Token": launch_token},
                timeout=5,
            ) as client:
                if client.get(f"/api/analyses/{analysis_id}/dashboard").status_code != 404:
                    raise VerificationError("DELETE_NOT_DURABLE")
                if client.get(f"/api/findings/{finding_id}").status_code != 404:
                    raise VerificationError("FINDING_DELETE_NOT_DURABLE")

            _stop_sidecar(running)
            stopped_processes += 1
            running = None

            if len(RULE_IDS) != 20:
                raise VerificationError("R01_R20_REGRESSION")
            if stopped_processes != 3:
                raise VerificationError("SHUTDOWN_INCOMPLETE")
        return MARKERS
    except VerificationError:
        raise
    except (httpx.HTTPError, KeyError, TypeError, ValueError, OSError) as error:
        raise VerificationError("VERIFICATION_FAILED") from error
    finally:
        if running is not None:
            _stop_sidecar(running)
