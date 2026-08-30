#!/usr/bin/env python3
from __future__ import annotations

import json
import os
import queue
import secrets
import subprocess
import sys
import tempfile
import threading
import time
from pathlib import Path

import httpx
from docx import Document

from qian_labor.rules.catalog import RULE_IDS

ROOT = Path(__file__).resolve().parents[1]
ENTRYPOINT = ROOT / "python" / "desktop_entrypoint.py"
READY_PREFIX = "QIAN_DESKTOP_READY="
TERMINAL = {"completed", "matching_review", "partial", "failed"}


def _write_fixture(path: Path) -> None:
    payload = {
        "synthetic_marker": "QIAN_DEMO_20260824",
        "document_type": "contract",
        "employee_number": "F-801",
        "employee_name": "完全虚构验收员工",
        "department": "虚构测试部门",
        "job_title": "虚构岗位",
        "facts": {
            "employment.status": "active",
            "employment.contract.exists": False,
            "employment.identity.match_status": "confirmed",
            "employment.material_coverage": 0.4,
            "analysis.minimum_core_coverage": 0.4,
        },
    }
    document = Document()
    document.add_heading("完全虚构桌面验收材料", level=1)
    document.add_paragraph(
        "QIAN_SYNTHETIC_JSON=" + json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
    )
    document.save(path)


def _start_sidecar(data_dir: Path, token: str) -> tuple[subprocess.Popen[str], str]:
    env = {
        **os.environ,
        "QIAN_DESKTOP_DATA_DIR": str(data_dir),
        "QIAN_DESKTOP_TOKEN": token,
        "QIAN_DESKTOP_PORT": "0",
    }
    process = subprocess.Popen(
        [sys.executable, str(ENTRYPOINT)],
        cwd=ROOT,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        bufsize=1,
    )
    assert process.stdout is not None
    lines: queue.Queue[str] = queue.Queue()

    def read_stdout() -> None:
        for line in process.stdout:
            lines.put(line.rstrip("\r\n"))

    threading.Thread(target=read_stdout, daemon=True).start()
    deadline = time.monotonic() + 15
    while time.monotonic() < deadline:
        if process.poll() is not None:
            raise RuntimeError("SIDECAR_EXITED_BEFORE_READY")
        try:
            line = lines.get(timeout=0.1)
        except queue.Empty:
            continue
        if not line.startswith(READY_PREFIX):
            continue
        payload = json.loads(line.split("=", 1)[1])
        if payload.get("host") != "127.0.0.1" or not isinstance(payload.get("port"), int):
            raise RuntimeError("SIDECAR_READY_INVALID")
        return process, f"http://127.0.0.1:{payload['port']}"
    raise RuntimeError("SIDECAR_READY_TIMEOUT")


def _stop_sidecar(process: subprocess.Popen[str]) -> None:
    if process.poll() is not None:
        return
    process.terminate()
    try:
        process.wait(timeout=5)
    except subprocess.TimeoutExpired:
        process.kill()
        process.wait(timeout=5)


def _poll(client: httpx.Client, analysis_id: str) -> dict[str, object]:
    deadline = time.monotonic() + 15
    while time.monotonic() < deadline:
        response = client.get(f"/api/analyses/{analysis_id}/processing")
        if response.status_code != 200:
            raise RuntimeError("PROCESSING_STATUS_FAILED")
        payload = response.json()
        if payload.get("status") in TERMINAL:
            return payload
        time.sleep(0.05)
    raise RuntimeError("PROCESSING_TIMEOUT")


def _assert_status(response: httpx.Response, expected: int, code: str) -> None:
    if response.status_code != expected:
        raise RuntimeError(code)


def main() -> int:
    process: subprocess.Popen[str] | None = None
    try:
        with tempfile.TemporaryDirectory(prefix="qian-desktop-verify-") as temp:
            temp_path = Path(temp)
            data_dir = temp_path / "app-data"
            fixture = temp_path / "fictional-contract.docx"
            _write_fixture(fixture)
            token = secrets.token_hex(32)

            process, base_url = _start_sidecar(data_dir, token)
            with httpx.Client(
                base_url=base_url,
                headers={"X-Qian-Desktop-Token": token},
                timeout=5,
            ) as client:
                _assert_status(client.get("/health"), 200, "HEALTH_FAILED")
                created = client.post(
                    "/api/analyses",
                    json={"name": "虚构桌面验收", "company_display_name": "完全虚构企业"},
                )
                _assert_status(created, 201, "ANALYSIS_CREATE_FAILED")
                analysis_id = str(created.json()["id"])
                imported = client.post(
                    f"/api/analyses/{analysis_id}/import-paths",
                    json={"paths": [str(fixture)]},
                )
                _assert_status(imported, 200, "SYNTHETIC_IMPORT_FAILED")
                submitted = client.post(f"/api/analyses/{analysis_id}/process")
                _assert_status(submitted, 202, "PROCESS_SUBMIT_FAILED")
                terminal = _poll(client, analysis_id)
                if terminal.get("status") != "completed":
                    raise RuntimeError("FAKE_PIPELINE_NOT_COMPLETED")
                dashboard = client.get(f"/api/analyses/{analysis_id}/dashboard")
                _assert_status(dashboard, 200, "DASHBOARD_FAILED")
                dashboard_payload = dashboard.json()
                findings = dashboard_payload.get("findings", [])
                if not findings:
                    raise RuntimeError("FINDING_MISSING")
                finding_id = str(findings[0]["id"])
                finding = client.get(f"/api/findings/{finding_id}")
                _assert_status(finding, 200, "FINDING_DETAIL_FAILED")
                if not finding.json().get("sources"):
                    raise RuntimeError("SOURCE_TRACE_MISSING")

            _stop_sidecar(process)
            process = None

            # A real process restart must preserve the completed analysis in the same SQLite file.
            process, base_url = _start_sidecar(data_dir, token)
            with httpx.Client(
                base_url=base_url,
                headers={"X-Qian-Desktop-Token": token},
                timeout=5,
            ) as client:
                persisted = client.get(f"/api/analyses/{analysis_id}/dashboard")
                _assert_status(persisted, 200, "SQLITE_PERSISTENCE_FAILED")
                deleted = client.delete(f"/api/analyses/{analysis_id}")
                _assert_status(deleted, 200, "DELETE_FAILED")
                if deleted.json().get("status") != "deleted":
                    raise RuntimeError("DELETE_STATUS_INVALID")

            _stop_sidecar(process)
            process = None

            # A second process restart proves deletion/tombstone durability.
            process, base_url = _start_sidecar(data_dir, token)
            with httpx.Client(
                base_url=base_url,
                headers={"X-Qian-Desktop-Token": token},
                timeout=5,
            ) as client:
                if client.get(f"/api/analyses/{analysis_id}/dashboard").status_code != 404:
                    raise RuntimeError("DELETE_NOT_DURABLE")
                if client.get(f"/api/findings/{finding_id}").status_code != 404:
                    raise RuntimeError("FINDING_DELETE_NOT_DURABLE")

            if len(RULE_IDS) != 20:
                raise RuntimeError("R01_R20_REGRESSION")

        for marker in (
            "SIDECAR_BOOT=PASS",
            "SQLITE_PERSISTENCE=PASS",
            "SYNTHETIC_IMPORT=PASS",
            "FAKE_PROVIDER_PIPELINE=PASS",
            "R01_R20_REGRESSION=PASS",
            "SOURCE_TRACE=PASS",
            "DELETE_CLEANUP=PASS",
        ):
            print(marker)
        return 0
    except Exception as error:  # process-level verifier prints only a safe type
        print(f"DESKTOP_VERIFY=FAIL:{type(error).__name__}", file=sys.stderr)
        return 1
    finally:
        if process is not None:
            _stop_sidecar(process)


if __name__ == "__main__":
    raise SystemExit(main())
